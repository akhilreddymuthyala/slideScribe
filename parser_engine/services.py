"""
Document Parser Agent
----------------------
Parses PDF, DOCX, README.md, and raw pasted text into structured JSON:
{
  "source_type": "pdf"|"docx"|"markdown"|"raw_text",
  "title": str,
  "sections": [{"heading": str, "level": int, "content": str}],
  "raw_text": str,
  "word_count": int,
  "char_count": int
}
"""

import re


class ParserError(Exception):
    pass


def clean_text(text: str) -> str:
    if not text:
        return ""
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'[^\S\n]+\n', '\n', text)
    return text.strip()


def guess_title(sections, fallback="Untitled Presentation"):
    for s in sections:
        if s.get('heading'):
            return s['heading']
    return fallback


def parse_pdf(file_obj) -> dict:
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise ParserError("PyMuPDF (fitz) is not installed.")

    try:
        file_bytes = file_obj.read()
        doc = fitz.open(stream=file_bytes, filetype="pdf")
    except Exception as e:
        raise ParserError(f"Could not open PDF: {e}")

    sections = []
    full_text_parts = []
    current_heading = None
    current_content = []

    for page in doc:
        blocks = page.get_text("dict")["blocks"]
        for block in blocks:
            if "lines" not in block:
                continue
            for line in block["lines"]:
                line_text = "".join(span["text"] for span in line["spans"]).strip()
                if not line_text:
                    continue
                max_size = max((span["size"] for span in line["spans"]), default=0)
                is_bold = any("Bold" in span.get("font", "") for span in line["spans"])
                is_heading = (max_size >= 13.5 and len(line_text) < 100) or (
                    is_bold and len(line_text) < 80 and max_size >= 11
                )

                if is_heading:
                    if current_heading or current_content:
                        sections.append({
                            "heading": current_heading or "Introduction",
                            "level": 1,
                            "content": clean_text(" ".join(current_content)),
                        })
                    current_heading = line_text
                    current_content = []
                else:
                    current_content.append(line_text)
                    full_text_parts.append(line_text)

    if current_heading or current_content:
        sections.append({
            "heading": current_heading or "Content",
            "level": 1,
            "content": clean_text(" ".join(current_content)),
        })

    raw_text = clean_text("\n".join(full_text_parts))
    doc.close()

    if not sections:
        sections = [{"heading": "Document Content", "level": 1, "content": raw_text}]

    return {
        "source_type": "pdf",
        "title": guess_title(sections),
        "sections": sections,
        "raw_text": raw_text,
        "word_count": len(raw_text.split()),
        "char_count": len(raw_text),
    }


def parse_docx(file_obj) -> dict:
    try:
        from docx import Document
    except ImportError:
        raise ParserError("python-docx is not installed.")

    try:
        doc = Document(file_obj)
    except Exception as e:
        raise ParserError(f"Could not open DOCX: {e}")

    sections = []
    full_text_parts = []
    current_heading = None
    current_content = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = (para.style.name or "").lower()
        is_heading = style_name.startswith("heading") or style_name == "title"

        if is_heading:
            if current_heading or current_content:
                sections.append({
                    "heading": current_heading or "Introduction",
                    "level": 1,
                    "content": clean_text(" ".join(current_content)),
                })
            current_heading = text
            current_content = []
        else:
            current_content.append(text)
            full_text_parts.append(text)

    if current_heading or current_content:
        sections.append({
            "heading": current_heading or "Content",
            "level": 1,
            "content": clean_text(" ".join(current_content)),
        })

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                full_text_parts.append(row_text)

    raw_text = clean_text("\n".join(full_text_parts))

    if not sections:
        sections = [{"heading": "Document Content", "level": 1, "content": raw_text}]

    return {
        "source_type": "docx",
        "title": guess_title(sections),
        "sections": sections,
        "raw_text": raw_text,
        "word_count": len(raw_text.split()),
        "char_count": len(raw_text),
    }


def parse_markdown(text: str) -> dict:
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    lines = text.split('\n')

    sections = []
    current_heading = None
    current_level = 1
    current_content = []
    full_text_parts = []

    heading_pattern = re.compile(r'^(#{1,6})\s+(.*)$')

    for line in lines:
        match = heading_pattern.match(line.strip())
        if match:
            if current_heading or current_content:
                sections.append({
                    "heading": current_heading or "Introduction",
                    "level": current_level,
                    "content": clean_text("\n".join(current_content)),
                })
            current_level = len(match.group(1))
            current_heading = match.group(2).strip()
            current_content = []
        else:
            stripped = line.strip()
            plain = re.sub(r'[`*_>#~\[\]]', '', stripped)
            plain = re.sub(r'\(.*?\)', '', plain).strip()
            if plain:
                current_content.append(stripped)
                full_text_parts.append(plain)

    if current_heading or current_content:
        sections.append({
            "heading": current_heading or "Content",
            "level": current_level,
            "content": clean_text("\n".join(current_content)),
        })

    raw_text = clean_text("\n".join(full_text_parts))

    if not sections:
        sections = [{"heading": "Document Content", "level": 1, "content": raw_text}]

    return {
        "source_type": "markdown",
        "title": guess_title(sections),
        "sections": sections,
        "raw_text": raw_text,
        "word_count": len(raw_text.split()),
        "char_count": len(raw_text),
    }


def parse_raw_text(text: str) -> dict:
    text = clean_text(text)
    paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]

    sections = []
    for i, para in enumerate(paragraphs):
        first_line = para.split('\n')[0]
        is_heading_like = len(first_line) < 60 and not first_line.endswith('.')
        if is_heading_like and len(para.split('\n')) > 1:
            heading = first_line
            content = '\n'.join(para.split('\n')[1:])
        else:
            heading = f"Section {i + 1}"
            content = para
        sections.append({"heading": heading, "level": 1, "content": clean_text(content)})

    if not sections:
        sections = [{"heading": "Notes", "level": 1, "content": text}]

    return {
        "source_type": "raw_text",
        "title": guess_title(sections, fallback="Untitled Notes"),
        "sections": sections,
        "raw_text": text,
        "word_count": len(text.split()),
        "char_count": len(text),
    }


EXTENSION_MAP = {
    'pdf': 'pdf', 'docx': 'docx',
    'md': 'markdown', 'markdown': 'markdown',
    'txt': 'raw_text',
}


def detect_type_from_filename(filename: str) -> str:
    ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
    return EXTENSION_MAP.get(ext, 'raw_text')


def parse_document(file_obj=None, filename: str = None, raw_text: str = None) -> dict:
    if raw_text is not None and file_obj is None:
        return parse_raw_text(raw_text)

    if file_obj is None or filename is None:
        raise ParserError("Either file_obj+filename or raw_text must be provided.")

    doc_type = detect_type_from_filename(filename)

    try:
        if doc_type == 'pdf':
            return parse_pdf(file_obj)
        elif doc_type == 'docx':
            return parse_docx(file_obj)
        elif doc_type == 'markdown':
            text = file_obj.read()
            if isinstance(text, bytes):
                text = text.decode('utf-8', errors='ignore')
            return parse_markdown(text)
        else:
            text = file_obj.read()
            if isinstance(text, bytes):
                text = text.decode('utf-8', errors='ignore')
            return parse_raw_text(text)
    except ParserError:
        raise
    except Exception as e:
        raise ParserError(f"Failed to parse {filename}: {e}")